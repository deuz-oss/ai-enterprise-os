"""Rendering dokumen generik (PDF) — infrastruktur bersama Fase 20
(Quotation item 2, Agreement item 3) dan Fase 21 item 4 (dokumen Job
Order). Satu fungsi generik `render_document_pdf` dipakai lintas jenis
dokumen lewat parameter (title/sections/footer/accent), bukan tiap jenis
dokumen bikin renderer sendiri -- pola sama seperti mesin auto-journal
(satu mesin, banyak event) alih-alih duplikasi logic.

Idiom reportlab (import lazy di dalam fungsi, buffer BytesIO, escape teks
user) sengaja disamakan dengan `talentpool/service.py::render_standard_cv`
supaya konsisten satu codebase, bukan pola baru.
"""

import io


def render_document_pdf(
    *,
    title: str,
    subtitle: str | None,
    sections: list[tuple[str, str]],
    footer_text: str | None = None,
    accent_color: str = "#0f172a",
) -> bytes:
    """Render PDF struktur sederhana: judul + subjudul + daftar field
    label/value + footer opsional. Cukup untuk Quotation/Agreement/dokumen
    JO -- bukan template builder drag-drop bebas (di luar cakupan v1, lihat
    `docs/02-product/PRD.md` Fase 20 item 2)."""
    from xml.sax.saxutils import escape

    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.platypus import HRFlowable, Paragraph, SimpleDocTemplate, Spacer

    accent = colors.HexColor(accent_color)
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        topMargin=20 * mm,
        bottomMargin=20 * mm,
        leftMargin=20 * mm,
        rightMargin=20 * mm,
        title=title,
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "DocTitle", parent=styles["Title"], fontSize=18, textColor=accent, spaceAfter=2
    )
    subtitle_style = ParagraphStyle(
        "DocSubtitle", parent=styles["Normal"], fontSize=10, textColor=colors.HexColor("#64748b")
    )
    label_style = ParagraphStyle(
        "FieldLabel", parent=styles["Normal"], fontSize=8.5, textColor=colors.HexColor("#64748b")
    )
    value_style = ParagraphStyle(
        "FieldValue", parent=styles["Normal"], fontSize=11, textColor=colors.HexColor("#0f172a")
    )

    def esc(value: str | None) -> str:
        return escape(value) if value is not None else ""

    story: list = [Paragraph(esc(title), title_style)]
    if subtitle:
        story.append(Paragraph(esc(subtitle), subtitle_style))
    story.append(Spacer(1, 3 * mm))
    story.append(HRFlowable(width="100%", color=accent, thickness=1))
    story.append(Spacer(1, 6 * mm))

    for label, value in sections:
        story.append(Paragraph(esc(label), label_style))
        story.append(Paragraph(esc(value) or "&#8212;", value_style))
        story.append(Spacer(1, 4 * mm))

    if footer_text:
        story.append(Spacer(1, 8 * mm))
        story.append(HRFlowable(width="100%", color=colors.HexColor("#e2e8f0"), thickness=0.5))
        story.append(Spacer(1, 2 * mm))
        story.append(Paragraph(esc(footer_text), label_style))

    doc.build(story)
    return buffer.getvalue()


def render_document_docx(
    *,
    title: str,
    subtitle: str | None,
    sections: list[tuple[str, str]],
    footer_text: str | None = None,
) -> bytes:
    """Render dokumen `.docx` struktur sederhana -- dipakai Agreement (item 3),
    yang PRD-nya minta output editable/`.docx` (beda dari Quotation yang
    cukup PDF). Belum ada precedent *authoring* `.docx` di codebase ini
    sebelum fungsi ini (python-docx sebelumnya cuma dipakai untuk *membaca*
    CV upload di talentpool) -- ground baru, bukan pola yang disalin."""
    import io

    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.add_heading(title, level=1)
    if subtitle:
        p = doc.add_paragraph(subtitle)
        if p.runs:
            p.runs[0].italic = True

    for label, value in sections:
        doc.add_heading(label, level=3)
        doc.add_paragraph(value or "-")

    if footer_text:
        doc.add_paragraph()
        footer = doc.add_paragraph(footer_text)
        if footer.runs:
            footer.runs[0].font.size = Pt(9)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def store_generated_document(
    *, object_prefix: str, file_name: str, data: bytes, content_type: str = "application/pdf"
) -> str:
    """Simpan dokumen hasil render ke object storage, kembalikan object_key --
    pola sama seperti `recruitment/service.py::_offering_letter_pdf`."""
    from app.core.storage import new_object_key, put_object

    object_key = new_object_key(object_prefix, file_name)
    put_object(object_key, data, content_type=content_type)
    return object_key
