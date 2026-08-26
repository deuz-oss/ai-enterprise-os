"""Talent Pool & CV Standardization (PRD §10) — Fase 13.

Pipeline: unggah (PDF/DOCX/scan/foto) → deteksi jenis dokumen → ekstraksi
(teks-layer atau satu panggilan vision LLM untuk scan) → validasi ketat ke
skema tetap + skor confidence per kelompok field → draft profil → review
recruiter (field ber-confidence rendah wajib dicek) → finalisasi yang
merender CV standar bertemplate branding tenant sebagai PDF berversi.

Prinsip: file asli tak pernah ditimpa; prompt & skema berversi sehingga
pipeline dapat dijalankan ulang.
"""

import io
import json
import logging
import re
from datetime import UTC, date, datetime

from app.core import storage
from app.core.database import parse_uuid
from app.core.llm import chat_completion, vision_completion
from app.modules import audit
from app.modules.recruitment.models import Candidate
from app.modules.talentpool.models import (
    CvDocKind,
    CvIntake,
    IntakeStatus,
    StandardCvVersion,
    TalentPoolStatus,
    TenantCvBranding,
)
from fastapi import HTTPException, UploadFile
from sqlalchemy import select
from sqlalchemy.orm import Session

logger = logging.getLogger(__name__)

ALLOWED_MIME = (
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "image/png",
    "image/jpeg",
    "image/webp",
)
MAX_CV_BYTES = 10 * 1024 * 1024

# Versi skema & prompt — naikkan saat mengubah struktur/prompt agar intake lama
# dapat ditandai untuk diproses ulang (PRD §10.1).
SCHEMA_VERSION = 1
PROMPT_VERSION = 1

# Ambang confidence; field di bawah ambang = wajib review manusia (PRD §10.1).
CONFIDENCE_THRESHOLD = 0.7

_EXTRACTION_PROMPT = (
    "Anda mesin ekstraksi CV Indonesia. Baca CV dan kembalikan HANYA JSON "
    "sesuai skema tetap berikut:\n"
    "{\n"
    '  "full_name": string,\n'
    '  "phone": string|null,\n'
    '  "email": string|null,\n'
    '  "domisili": string|null,\n'
    '  "birth_date": "YYYY-MM-DD"|null,\n'
    '  "summary": string|null,\n'
    '  "education": [{"jenjang": string, "institusi": string, '
    '"jurusan": string|null, "tahun_lulus": number|null, "ipk": number|null}],\n'
    '  "experience": [{"perusahaan": string, "posisi": string, "periode": string, '
    '"industri": string|null, "ringkasan": string|null}],\n'
    '  "skills": [string],\n'
    '  "certifications": [{"nama": string, "penerbit": string|null, "tahun": number|null}],\n'
    '  "languages": [{"bahasa": string, "tingkat": string}],\n'
    '  "readiness": "segera"|"n_minggu"|"belum_tentu",\n'
    '  "readiness_weeks": number|null,\n'
    '  "willing_locations": [string],\n'
    '  "expected_salary": number|null,\n'
    '  "contract_preference": string|null,\n'
    '  "confidence": {"identitas": 0..1, "pendidikan": 0..1, "pengalaman": 0..1, '
    '"skill": 0..1, "penempatan": 0..1}\n'
    "}\n"
    "Gunakan null/[] untuk data yang tidak ada. Jangan mengarang angka atau "
    "tanggal. Nilai confidence mencerminkan seberapa jelas data terbaca."
)

_EMAIL_RE = re.compile(r"^[\w.+-]+@[\w-]+\.[\w.-]+$")
_PHONE_RE = re.compile(r"^\+?[\d()\-\s]{8,20}$")
_DATE_RE = re.compile(r"^\d{4}-\d{2}-\d{2}$")


# ---------- Deteksi jenis dokumen & ekstraksi teks ----------


def detect_doc_kind(data: bytes, file_name: str, mime_type: str) -> CvDocKind:
    image_exts = (".png", ".jpg", ".jpeg", ".webp")
    if mime_type.startswith("image/") or file_name.lower().endswith(image_exts):
        return CvDocKind.image
    if file_name.lower().endswith(".docx") or mime_type.endswith("wordprocessingml.document"):
        return CvDocKind.docx
    if file_name.lower().endswith(".pdf") or data[:5] == b"%PDF-":
        return CvDocKind.pdf_text if _pdf_has_text_layer(data) else CvDocKind.pdf_scan
    raise HTTPException(status_code=422, detail="Format CV tidak didukung (PDF/DOCX/gambar)")


def _pdf_has_text_layer(data: bytes) -> bool:
    from pypdf import PdfReader

    try:
        reader = PdfReader(io.BytesIO(data))
        text = "\n".join((page.extract_text() or "") for page in reader.pages[:3])
    except Exception as exc:  # noqa: BLE001
        raise HTTPException(
            status_code=422, detail="PDF gagal dibaca; pastikan file tidak rusak"
        ) from exc
    return len(text.strip()) >= 40


def _docx_text(data: bytes) -> str:
    import docx  # python-docx

    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as exc:  # noqa: BLE001
        raise HTTPException(status_code=422, detail="Dokumen DOCX gagal dibaca") from exc
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text.strip() for cell in row.cells))
    return "\n".join(parts)


def extract_profile(db: Session, data: bytes, kind: CvDocKind) -> dict:
    """Satu panggilan LLM: teks-layer → json_mode; scan/foto → vision (§10.4)."""
    if kind == CvDocKind.pdf_scan or kind == CvDocKind.image:
        import base64

        mime = "image/png" if kind == CvDocKind.image else "application/pdf"
        raw = vision_completion(
            _EXTRACTION_PROMPT,
            "Ekstrak data CV dari dokumen hasil scan ini.",
            image_b64=base64.b64encode(data).decode(),
            mime_type=mime,
        )
    else:
        text = _docx_text(data) if kind == CvDocKind.docx else _pdf_text(data)
        raw = chat_completion(_EXTRACTION_PROMPT, f"CV:\n{text[:24000]}")
    return raw if isinstance(raw, dict) else {}


def _pdf_text(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    return "\n".join((page.extract_text() or "") for page in reader.pages)


# ---------- Validasi ketat + confidence (deterministik) ----------


def normalize_and_score(raw: dict) -> tuple[dict, dict, list[str]]:
    """Normalisasi ke skema tetap; hitung confidence per kelompok; daftar field
    yang wajib direview (confidence < ambang atau gagal validasi)."""
    conf_raw = raw.get("confidence")
    conf_in: dict = conf_raw if isinstance(conf_raw, dict) else {}
    groups = {
        "identitas": float(conf_in.get("identitas") or 0),
        "pendidikan": float(conf_in.get("pendidikan") or 0),
        "pengalaman": float(conf_in.get("pengalaman") or 0),
        "skill": float(conf_in.get("skill") or 0),
        "penempatan": float(conf_in.get("penempatan") or 0),
    }

    full_name = str(raw.get("full_name") or "").strip()
    phone = str(raw.get("phone") or "").strip() or None
    email = str(raw.get("email") or "").strip().lower() or None
    domisili = str(raw.get("domisili") or "").strip() or None
    birth_date = str(raw.get("birth_date") or "").strip() or None

    # Koreksi deterministik atas confidence model.
    if full_name:
        groups["identitas"] = max(groups["identitas"], 0.6)
    if email and _EMAIL_RE.match(email):
        groups["identitas"] = min(1.0, groups["identitas"] + 0.15)
    elif email:
        groups["identitas"] -= 0.2
        email = None
    if phone and not _PHONE_RE.match(phone):
        groups["identitas"] -= 0.2
        phone = None
    if birth_date and not _DATE_RE.match(birth_date):
        birth_date = None
    groups["identitas"] = min(max(groups["identitas"], 0.0), 1.0)

    education = _list_of_dicts(raw.get("education"))
    experience = _list_of_dicts(raw.get("experience"))
    certifications = [c for c in education_and_certs(raw) if c]  # sertifikasi divalidasi di helper
    skills = [str(s).strip()[:100] for s in _list_of_str(raw.get("skills"))][:60]
    languages = _list_of_dicts(raw.get("languages"))

    if not education:
        groups["pendidikan"] = min(groups["pendidikan"], 0.5)
    if not experience:
        groups["pengalaman"] = min(groups["pengalaman"], 0.5)
    if not skills:
        groups["skill"] = min(groups["skill"], 0.4)

    readiness_raw = str(raw.get("readiness") or "").strip()
    readiness = readiness_raw if readiness_raw in ("segera", "n_minggu", "belum_tentu") else None
    try:
        readiness_weeks_raw = raw.get("readiness_weeks")
        readiness_weeks = int(readiness_weeks_raw) if readiness_weeks_raw is not None else None
    except (TypeError, ValueError):
        readiness_weeks = None
    willing_locations = [str(x).strip()[:120] for x in _list_of_str(raw.get("willing_locations"))]
    try:
        expected_salary_raw = raw.get("expected_salary")
        expected_salary = round(float(expected_salary_raw)) if expected_salary_raw else None
    except (TypeError, ValueError):
        expected_salary = None
    contract_pref = str(raw.get("contract_preference") or "").strip()[:120] or None

    if not readiness:
        groups["penempatan"] = min(groups["penempatan"], 0.5)
    groups = {k: round(min(max(v, 0.0), 1.0), 2) for k, v in groups.items()}

    profile: dict = {
        "full_name": full_name[:255] or None,
        "phone": phone,
        "email": email,
        "domisili": domisili,
        "birth_date": birth_date,
        "summary": str(raw.get("summary") or "").strip()[:2000] or None,
        "education": education,
        "experience": experience,
        "certifications": certifications,
        "skills": skills,
        "languages": languages,
        "readiness": readiness,
        "readiness_weeks": readiness_weeks,
        "willing_locations": willing_locations,
        "expected_salary": expected_salary,
        "contract_preference": contract_pref,
    }

    needs_review = sorted(k for k, v in groups.items() if v < CONFIDENCE_THRESHOLD)
    if not profile["full_name"]:
        needs_review = sorted(set(needs_review) | {"identitas"})
    return profile, groups, needs_review


def education_and_certs(raw: dict) -> list[dict]:
    certs: list[dict] = []
    for item in _list_of_dicts(raw.get("certifications")):
        nama = str(item.get("nama") or "").strip()
        if not nama:
            continue
        tahun_raw = item.get("tahun")
        try:
            tahun = int(tahun_raw) if tahun_raw is not None else None
        except (TypeError, ValueError):
            tahun = None
        certs.append(
            {
                "nama": nama[:255],
                "penerbit": str(item.get("penerbit") or "").strip()[:255] or None,
                "tahun": tahun,
            }
        )
    return certs


def _list_of_str(value) -> list[str]:
    if not isinstance(value, list):
        return []
    return [str(v) for v in value if str(v).strip()]


def _list_of_dicts(value) -> list[dict]:
    if not isinstance(value, list):
        return []
    out: list[dict] = []
    for item in value[:50]:
        if isinstance(item, dict):
            out.append({k: v for k, v in item.items() if v not in (None, "")})
    return out


def _safe_date(value) -> str | None:
    try:
        return date.fromisoformat(str(value).strip()).isoformat()
    except (TypeError, ValueError):
        return None


# ---------- Intake / review / finalisasi ----------


async def intake_cv(
    db: Session,
    *,
    user,
    file: UploadFile,
    candidate_id: str | None = None,
    consent: bool = False,
) -> CvIntake:
    """Unggah CV → simpan original → jalankan pipeline ekstraksi.

    Bila AI tidak tersedia/ekstraksi gagal: intake tetap tersimpan dengan
    status `gagal` dan bisa diproses ulang nanti (PRD §10.1 re-process).
    """
    from app.modules.recruitment.service import _get_candidate

    if not consent:
        raise HTTPException(
            status_code=422,
            detail="Persetujuan pemrosesan data pribadi (UU PDP) wajib dicentang",
        )
    mime = (file.content_type or "").lower()
    if mime not in ALLOWED_MIME:
        raise HTTPException(status_code=422, detail="Format CV harus PDF, DOCX, atau gambar")
    data = await file.read()
    if len(data) > MAX_CV_BYTES:
        raise HTTPException(status_code=422, detail="Ukuran CV maksimal 10 MB")

    candidate = (
        _get_candidate(db, candidate_id)
        if candidate_id
        else Candidate(full_name=f"Kandidat baru {date.today().isoformat()}")
    )
    db.add(candidate)
    db.flush()

    object_key = storage.new_object_key(f"talentpool/{candidate.id}", file.filename or "cv.pdf")
    storage.put_object(object_key, data, mime)

    intake = CvIntake(
        tenant_id=candidate.tenant_id,
        candidate_id=candidate.id,
        uploaded_by_id=user.id,
        file_name=(file.filename or "cv.pdf")[:255],
        mime_type=mime,
        object_key=object_key,
        file_size=len(data),
        schema_version=SCHEMA_VERSION,
        prompt_version=PROMPT_VERSION,
        consent=True,
        status=IntakeStatus.processing,
    )
    db.add(intake)
    db.commit()
    audit.log_event(
        db,
        action="talentpool.cv_uploaded",
        entity_type="cv_intake",
        entity_id=intake.id,
        object_key=object_key,
        detail={"candidate": str(candidate.id), "file_name": intake.file_name},
    )

    _process_intake(db, intake, data)
    return intake


def _process_intake(db: Session, intake: CvIntake, data: bytes | None = None) -> None:
    """Jalankan ekstraksi untuk satu intake; gagal ditandai, tidak melempar."""
    if data is None:
        data = storage.get_object(intake.object_key)
    try:
        kind = detect_doc_kind(data, intake.file_name, intake.mime_type)
        raw = extract_profile(db, data, kind)
        profile, confidences, needs_review = normalize_and_score(raw)
    except Exception as exc:  # noqa: BLE001 - kegagalan AI tidak boleh menggugurkan unggahan
        intake.status = IntakeStatus.failed
        intake.error = str(getattr(exc, "detail", exc))[:500]
        db.commit()
        return

    intake.doc_kind = kind
    intake.extracted = json.dumps(profile, ensure_ascii=False)
    intake.confidences = json.dumps(confidences)
    intake.needs_review = json.dumps(needs_review)
    intake.reviewed_fields = json.dumps([])
    intake.readiness = profile.get("readiness")
    intake.tp_status = TalentPoolStatus.diproses
    intake.status = IntakeStatus.review
    intake.error = None
    intake.processed_at = datetime.now(UTC)

    candidate = db.get(Candidate, intake.candidate_id)
    if candidate is not None and profile.get("full_name"):
        _sync_candidate(candidate, profile)
    # Intake lama non-final pada kandidat yang sama diarsipkan statusnya.
    db.commit()


def _sync_candidate(candidate: Candidate, p: dict) -> None:
    candidate.full_name = p.get("full_name") or candidate.full_name
    candidate.phone = p.get("phone") or candidate.phone
    candidate.email = p.get("email") or candidate.email
    candidate.city = p.get("domisili") or candidate.city
    candidate.expected_salary = p.get("expected_salary") or candidate.expected_salary
    if p.get("skills"):
        candidate.skills = ", ".join(p["skills"])
    if p.get("experience"):
        candidate.current_company = p["experience"][0].get("perusahaan")


def get_intake(db: Session, intake_id: str) -> dict:
    intake = db.get(CvIntake, parse_uuid(intake_id))
    if intake is None:
        raise HTTPException(status_code=404, detail="Intake CV tidak ditemukan")
    return serialize_intake(db, intake)


def serialize_intake(db: Session, intake: CvIntake) -> dict:
    candidate = db.get(Candidate, intake.candidate_id)
    versions = (
        db.execute(
            select(StandardCvVersion)
            .where(StandardCvVersion.candidate_id == intake.candidate_id)
            .order_by(StandardCvVersion.seq.desc())
        )
        .scalars()
        .all()
    )
    return {
        "id": str(intake.id),
        "candidate_id": str(intake.candidate_id),
        "candidate_name": candidate.full_name if candidate else None,
        "status": intake.status.value,
        "doc_kind": intake.doc_kind.value if intake.doc_kind else None,
        "file_name": intake.file_name,
        "extracted": json.loads(intake.extracted) if intake.extracted else None,
        "confidences": json.loads(intake.confidences) if intake.confidences else {},
        "needs_review": json.loads(intake.needs_review) if intake.needs_review else [],
        "reviewed_fields": json.loads(intake.reviewed_fields) if intake.reviewed_fields else [],
        "schema_version": intake.schema_version,
        "prompt_version": intake.prompt_version,
        "readiness": intake.readiness,
        "tp_status": intake.tp_status.value,
        "error": intake.error,
        "versions": [
            {
                "id": str(v.id),
                "seq": v.seq,
                "is_locked": v.is_locked,
                "created_at": v.created_at.isoformat(),
                "download_url": f"/api/v1/talentpool/cv-versions/{v.id}/download",
            }
            for v in versions
        ],
    }


def review_intake(db: Session, *, user, intake_id: str, corrections: dict, reviewed: list[str]):
    """Koreksi recruiter atas draft profil + tandai kelompok field sudah dicek."""
    intake = db.get(CvIntake, parse_uuid(intake_id))
    if intake is None:
        raise HTTPException(status_code=404, detail="Intake CV tidak ditemukan")
    if intake.status != IntakeStatus.review:
        raise HTTPException(status_code=409, detail="Intake tidak dalam tahap review")

    profile = json.loads(intake.extracted) if intake.extracted else {}
    allowed_top = set(profile.keys())
    for key, value in (corrections or {}).items():
        if key in allowed_top:
            profile[key] = value
    groups, needs_review = rescore(profile)
    intake.extracted = json.dumps(profile, ensure_ascii=False)
    intake.confidences = json.dumps(groups)
    already = set(json.loads(intake.reviewed_fields or "[]"))
    already |= {str(r) for r in (reviewed or [])}
    intake.reviewed_fields = json.dumps(sorted(already & set(groups.keys())))
    intake.needs_review = json.dumps(needs_review)
    db.commit()
    audit.log_event(
        db,
        action="talentpool.reviewed",
        entity_type="cv_intake",
        entity_id=intake.id,
        detail={"reviewed": sorted(already)},
    )
    return intake


def rescore(profile: dict) -> tuple[dict, list[str]]:
    """Skor ulang deterministik setelah koreksi manual — koreksi = pasti benar."""
    groups = {
        "identitas": 0.9,
        "pendidikan": 0.9,
        "pengalaman": 0.9,
        "skill": 0.9,
        "penempatan": 0.9,
    }
    p = profile
    if not p.get("full_name") or not p.get("domisili"):
        groups["identitas"] = 0.5
    if not p.get("education"):
        groups["pendidikan"] = 0.5
    if not p.get("experience"):
        groups["pengalaman"] = 0.5
    if not p.get("skills"):
        groups["skill"] = 0.5
    if not p.get("readiness"):
        groups["penempatan"] = 0.5
    needs_review = sorted(k for k, v in groups.items() if v < CONFIDENCE_THRESHOLD)
    return groups, needs_review


def finalize_intake(db: Session, *, user, intake_id: str):
    """Finalisasi: semua field wajib-review harus sudah dicek → render CV standar.

    Membuat versi baru dokumen CV standar (snapshot PDF berversi, §10.3).
    """
    intake = db.get(CvIntake, parse_uuid(intake_id))
    if intake is None:
        raise HTTPException(status_code=404, detail="Intake CV tidak ditemukan")
    if intake.status != IntakeStatus.review:
        raise HTTPException(
            status_code=409, detail="Hanya intake berstatus review yang difinalisasi"
        )
    needs = set(json.loads(intake.needs_review or "[]"))
    reviewed = set(json.loads(intake.reviewed_fields or "[]"))
    blocking = sorted(needs - reviewed)
    if blocking:
        raise HTTPException(
            status_code=422,
            detail=f"Kelompok field berikut wajib direview dahulu: {', '.join(blocking)}",
        )

    profile = json.loads(intake.extracted or "{}")
    branding = get_branding(db)
    photo_bytes = None
    if branding.show_photo:
        candidate_row = db.get(Candidate, intake.candidate_id)
        if candidate_row is not None:
            photo_bytes = candidate_photo_bytes(db, candidate_row)
    pdf_bytes = render_standard_cv(db, profile, branding, photo_bytes=photo_bytes)
    version_key = storage.new_object_key(
        f"talentpool/{intake.candidate_id}/standard-cv",
        f"cv-standar-v{str(intake.id)[:8]}.pdf",
    )
    storage.put_object(version_key, pdf_bytes, "application/pdf")

    last_seq = db.scalar(
        select(StandardCvVersion.seq)
        .where(StandardCvVersion.candidate_id == intake.candidate_id)
        .order_by(StandardCvVersion.seq.desc())
        .limit(1)
    )
    version = StandardCvVersion(
        tenant_id=intake.tenant_id,
        candidate_id=intake.candidate_id,
        intake_id=intake.id,
        seq=(last_seq or 0) + 1,
        object_key=version_key,
        file_size=len(pdf_bytes),
        created_by_id=user.id,
    )
    db.add(version)
    intake.status = IntakeStatus.finalized
    intake.tp_status = TalentPoolStatus.diproses
    db.commit()
    db.refresh(version)
    audit.log_event(
        db,
        action="talentpool.finalized",
        entity_type="cv_intake",
        entity_id=intake.id,
        object_key=version_key,
        detail={"candidate": str(intake.candidate_id), "version_seq": version.seq},
    )
    return intake


def reprocess_intake(db: Session, *, intake_id: str) -> CvIntake:
    """Jalankan ulang pipeline dengan skema/prompt terkini (PRD §10.1)."""
    intake = db.get(CvIntake, parse_uuid(intake_id))
    if intake is None:
        raise HTTPException(status_code=404, detail="Intake CV tidak ditemukan")
    if intake.status == IntakeStatus.finalized:
        raise HTTPException(
            status_code=409,
            detail="Intake sudah finalisasi — unggah CV baru untuk versi baru",
        )
    intake.schema_version = SCHEMA_VERSION
    intake.prompt_version = PROMPT_VERSION
    intake.status = IntakeStatus.processing
    db.commit()
    _process_intake(db, intake)
    return intake


def download_version(db: Session, version_id: str) -> tuple[bytes, str]:
    version = db.get(StandardCvVersion, parse_uuid(version_id))
    if version is None:
        raise HTTPException(status_code=404, detail="Versi CV tidak ditemukan")
    data = storage.get_object(version.object_key)
    name = f"cv-standar-{version.seq:03d}.pdf"
    audit.log_event(
        db,
        action="talentpool.cv_version_downloaded",
        entity_type="standard_cv_version",
        entity_id=version.id,
        object_key=version.object_key,
    )
    return data, name


def lock_version_for_placement(db: Session, *, candidate_id, placement_id) -> None:
    """Snapshot submission (§10.3): kunci versi CV terbaru saat kandidat diusulkan."""
    version = (
        db.execute(
            select(StandardCvVersion)
            .where(
                StandardCvVersion.candidate_id == parse_uuid(str(candidate_id)),
                StandardCvVersion.is_locked.is_(False),
            )
            .order_by(StandardCvVersion.seq.desc())
            .limit(1)
        )
        .scalars()
        .first()
    )
    if version is None:
        return
    version.is_locked = True
    version.locked_for_placement_id = parse_uuid(str(placement_id))
    db.commit()


# ---------- Talent pool listing & hak hapus ----------


def list_talentpool(
    db: Session,
    *,
    q: str | None = None,
    domisili: str | None = None,
    skill: str | None = None,
    readiness: str | None = None,
    tp_status: str | None = None,
    has_standard_cv: bool | None = None,
) -> list[dict]:
    stmt = select(Candidate).order_by(Candidate.created_at.desc())
    if q:
        stmt = stmt.where(Candidate.full_name.ilike(f"%{q}%"))
    if domisili:
        stmt = stmt.where(Candidate.city.ilike(f"%{domisili}%"))
    if skill:
        stmt = stmt.where(Candidate.skills.ilike(f"%{skill}%"))
    candidates = list(db.execute(stmt).scalars())

    intakes = {
        i.candidate_id: i
        for i in db.execute(select(CvIntake).order_by(CvIntake.created_at.desc())).scalars()
    }
    locked_counts: dict = {}
    for v in db.execute(select(StandardCvVersion)).scalars():
        locked_counts[v.candidate_id] = max(locked_counts.get(v.candidate_id, 0), v.seq)

    rows: list[dict] = []
    for c in candidates:
        intake = intakes.get(c.id)
        if readiness and (not intake or intake.readiness != readiness):
            continue
        if tp_status and (not intake or intake.tp_status.value != tp_status):
            continue
        if has_standard_cv and c.id not in locked_counts:
            continue
        rows.append(
            {
                "candidate_id": str(c.id),
                "full_name": c.full_name,
                "city": c.city,
                "email": c.email,
                "phone": c.phone,
                "expected_salary": float(c.expected_salary) if c.expected_salary else None,
                "skills": c.skills,
                "readiness": intake.readiness if intake else None,
                "tp_status": intake.tp_status.value if intake else TalentPoolStatus.baru.value,
                "intake_status": intake.status.value if intake else None,
                "latest_intake_id": str(intake.id) if intake else None,
                "needs_review_count": (
                    len(json.loads(intake.needs_review)) if intake and intake.needs_review else 0
                ),
                "latest_cv_version": locked_counts.get(c.id),
            }
        )
    return rows


def forget_candidate(db: Session, *, user, candidate_id: str) -> dict:
    """Hak hapus subjek data (UU PDP, PRD §10.5): hapus profil & snapshot,
    scrub PII kandidat. Jejak audit tetap tanpa data pribadi."""
    candidate = db.get(Candidate, parse_uuid(candidate_id))
    if candidate is None:
        raise HTTPException(status_code=404, detail="Kandidat tidak ditemukan")
    intakes = (
        db.execute(select(CvIntake).where(CvIntake.candidate_id == candidate.id)).scalars().all()
    )
    removed = {"intakes": len(intakes), "versions": 0}
    for i in intakes:
        db.delete(i)
    versions = (
        db.execute(select(StandardCvVersion).where(StandardCvVersion.candidate_id == candidate.id))
        .scalars()
        .all()
    )
    removed["versions"] = len(versions)
    for v in versions:
        db.delete(v)
    candidate.full_name = "(dihapus atas permintaan)"
    candidate.phone = None
    candidate.email = None
    candidate.city = None
    candidate.skills = None
    candidate.notes = "Data pribadi dihapus permintaan subjek (UU PDP)"
    db.commit()
    audit.log_event(
        db,
        action="talentpool.forgotten",
        entity_type="candidate",
        entity_id=candidate.id,
        detail=removed,
    )
    return removed


ALLOWED_PHOTO_MIME = ("image/png", "image/jpeg")
_MAX_PHOTO_BYTES = 5 * 1024 * 1024


def upload_candidate_photo(db: Session, *, user, candidate_id: str, data: bytes, mime: str):
    """Unggah foto kandidat untuk CV standar (PNG/JPEG ≤ 5 MB)."""
    from app.modules.recruitment.service import _get_candidate

    if mime not in ALLOWED_PHOTO_MIME:
        raise HTTPException(status_code=422, detail="Foto harus PNG atau JPEG")
    if not data:
        raise HTTPException(status_code=422, detail="File foto kosong")
    if len(data) > _MAX_PHOTO_BYTES:
        raise HTTPException(status_code=422, detail="Foto maksimal 5 MB")
    candidate = _get_candidate(db, candidate_id)
    key = storage.new_object_key(
        f"talentpool/{candidate.id}/photo", "photo.png" if mime == "image/png" else "photo.jpg"
    )
    storage.put_object(key, data, mime)
    candidate.photo_object_key = key
    db.commit()
    audit.log_event(
        db,
        action="talentpool.photo_uploaded",
        entity_type="candidate",
        entity_id=candidate.id,
        object_key=key,
        detail={"by": getattr(user, "email", "?"), "size": len(data)},
    )
    return {"candidate_id": str(candidate.id), "has_photo": True}


def remove_candidate_photo(db: Session, *, user, candidate_id: str):
    from app.modules.recruitment.service import _get_candidate

    candidate = _get_candidate(db, candidate_id)
    candidate.photo_object_key = None
    db.commit()
    audit.log_event(
        db,
        action="talentpool.photo_removed",
        entity_type="candidate",
        entity_id=candidate.id,
        detail={"by": getattr(user, "email", "?")},
    )
    return {"candidate_id": str(candidate.id), "has_photo": False}


def candidate_photo_bytes(db: Session, candidate) -> bytes | None:
    if not candidate.photo_object_key:
        return None
    try:
        return storage.get_object(candidate.photo_object_key)
    except Exception:  # noqa: BLE001 - foto hilang tidak boleh gagalkan render
        logger.warning("Foto kandidat %s gagal dibaca", candidate.id)
        return None


# ---------- Branding per tenant ----------


def get_branding(db: Session) -> TenantCvBranding:
    branding = db.execute(select(TenantCvBranding)).scalars().first()
    if branding is None:
        branding = TenantCvBranding(tenant_id=None)  # tenant diisi listener
        db.add(branding)
        db.commit()
        db.refresh(branding)
    return branding


def update_branding(db: Session, *, user, payload: dict) -> TenantCvBranding:
    branding = get_branding(db)
    accent = str(payload.get("accent_color") or "").strip()
    if accent:
        if not re.match(r"^#[0-9a-fA-F]{6}$", accent):
            raise HTTPException(status_code=422, detail="accent_color harus format #RRGGBB")
        branding.accent_color = accent
    if "footer_text" in payload:
        branding.footer_text = str(payload.get("footer_text") or "").strip()[:255]
    if "show_photo" in payload:
        branding.show_photo = bool(payload.get("show_photo"))
    db.commit()
    audit.log_event(
        db,
        action="talentpool.branding_updated",
        entity_type="tenant_cv_branding",
        entity_id=branding.id,
        detail={"by": getattr(user, "email", "?")},
    )
    return branding


ALLOWED_LOGO_MIME = ("image/png", "image/jpeg")
_MAX_LOGO_BYTES = 2 * 1024 * 1024


def upload_branding_logo(db: Session, *, user, data: bytes, mime: str) -> TenantCvBranding:
    """Unggah logo perusahaan untuk header CV standar (PNG/JPEG ≤ 2 MB)."""
    branding = get_branding(db)
    if mime not in ALLOWED_LOGO_MIME:
        raise HTTPException(status_code=422, detail="Logo harus PNG atau JPEG")
    if len(data) == 0:
        raise HTTPException(status_code=422, detail="File logo kosong")
    if len(data) > _MAX_LOGO_BYTES:
        raise HTTPException(status_code=422, detail="Logo maksimal 2 MB")
    key = storage.new_object_key("branding", "logo.png" if mime == "image/png" else "logo.jpg")
    storage.put_object(key, data, mime)
    branding.logo_object_key = key
    db.commit()
    audit.log_event(
        db,
        action="talentpool.logo_uploaded",
        entity_type="tenant_cv_branding",
        entity_id=branding.id,
        object_key=key,
        detail={"by": getattr(user, "email", "?"), "size": len(data)},
    )
    return branding


def remove_branding_logo(db: Session, *, user) -> TenantCvBranding:
    branding = get_branding(db)
    branding.logo_object_key = None
    db.commit()
    audit.log_event(
        db,
        action="talentpool.logo_removed",
        entity_type="tenant_cv_branding",
        entity_id=branding.id,
        detail={"by": getattr(user, "email", "?")},
    )
    return branding


def serialize_branding(branding: TenantCvBranding) -> dict:
    return {
        "accent_color": branding.accent_color,
        "footer_text": branding.footer_text,
        "show_photo": branding.show_photo,
        "has_logo": branding.logo_object_key is not None,
        "logo_url": (
            "/api/v1/talentpool/branding/logo/download" if branding.logo_object_key else None
        ),
    }


# ---------- Render CV standar (reportlab, §10.3) ----------


def render_standard_cv(
    db: Session, profile: dict, branding: TenantCvBranding, photo_bytes: bytes | None = None
) -> bytes:
    """PDF CV standar struktur tetap: identitas → ringkasan → pengalaman →
    pendidikan → skill/sertifikasi/bahasa → data penempatan.

    Foto kandidat tampil di kanan header bila branding tenant mengizinkan
    (show_photo) dan fotonya tersedia.
    """
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.platypus import (
        HRFlowable,
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )

    accent = colors.HexColor(branding.accent_color)
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        topMargin=16 * mm,
        bottomMargin=18 * mm,
        leftMargin=18 * mm,
        rightMargin=18 * mm,
        title="CV Standar",
    )
    styles = getSampleStyleSheet()
    h1 = ParagraphStyle(
        "CVName", parent=styles["Title"], fontSize=20, textColor=accent, spaceAfter=2
    )
    sub = ParagraphStyle(
        "CVSub", parent=styles["Normal"], fontSize=9.5, textColor=colors.HexColor("#555555")
    )
    section = ParagraphStyle(
        "CVSection",
        parent=styles["Heading2"],
        fontSize=12,
        textColor=accent,
        spaceBefore=8,
        spaceAfter=3,
    )
    body = ParagraphStyle("CVBody", parent=styles["Normal"], fontSize=9.5, leading=13)
    small = ParagraphStyle(
        "CVSmall", parent=styles["Normal"], fontSize=8.5, textColor=colors.HexColor("#555555")
    )

    story: list = []

    def esc(value) -> str:
        from xml.sax.saxutils import escape

        return escape(str(value)) if value is not None else ""

    # Logo tenant di atas header bila tersedia (tinggi tetap, lebar mengikuti rasio).
    if branding.logo_object_key:
        try:
            from reportlab.lib.utils import ImageReader
            from reportlab.platypus import Image

            logo_bytes = storage.get_object(branding.logo_object_key)
            reader = ImageReader(io.BytesIO(logo_bytes))
            iw, ih = reader.getSize()
            height = 14 * mm
            width = min(height * (iw / max(ih, 1)), 45 * mm)
            story.append(Image(reader, width=width, height=height, hAlign="LEFT"))
            story.append(Spacer(1, 3 * mm))
        except Exception:  # noqa: BLE001 - logo rusak tidak boleh gagalkan render
            pass

    name_para = Paragraph(esc(profile.get("full_name") or "-"), h1)
    contact_bits = [
        b for b in (profile.get("phone"), profile.get("email"), profile.get("domisili")) if b
    ]
    contact_para = Paragraph(" · ".join(esc(b) for b in contact_bits), sub)

    photo_img = None
    if photo_bytes:
        try:
            from reportlab.lib.utils import ImageReader
            from reportlab.platypus import Image

            reader = ImageReader(io.BytesIO(photo_bytes))
            iw, ih = reader.getSize()
            height = 24 * mm
            width = min(height * (iw / max(ih, 1)), 20 * mm)
            photo_img = Image(reader, width=width, height=height)
        except Exception:  # noqa: BLE001
            photo_img = None

    if photo_img is not None:
        header_tbl = Table(
            [[name_para, photo_img], [contact_para, ""]],
            colWidths=[120 * mm, 30 * mm],
        )
        header_tbl.setStyle(
            TableStyle(
                [
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("ALIGN", (1, 0), (1, -1), "RIGHT"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 0),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                    ("TOPPADDING", (0, 0), (-1, -1), 0),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 1),
                ]
            )
        )
        story.append(header_tbl)
    else:
        story.append(name_para)
        story.append(contact_para)
    if profile.get("birth_date"):
        story.append(Paragraph(f"Tanggal lahir: {esc(profile['birth_date'])}", small))
    story.append(HRFlowable(width="100%", thickness=1.2, color=accent, spaceBefore=6, spaceAfter=6))

    if profile.get("summary"):
        story.append(Paragraph("Ringkasan Profil", section))
        story.append(Paragraph(esc(profile["summary"]), body))

    if profile.get("experience"):
        story.append(Paragraph("Pengalaman Kerja", section))
        for exp in profile["experience"]:
            head = f"<b>{esc(exp.get('posisi'))}</b> — {esc(exp.get('perusahaan'))}"
            if exp.get("periode"):
                head += f" ({esc(exp['periode'])})"
            story.append(Paragraph(head, body))
            if exp.get("ringkasan"):
                story.append(Paragraph(esc(exp["ringkasan"]), small))
            if exp.get("industri"):
                story.append(Paragraph(f"Industri: {esc(exp['industri'])}", small))

    if profile.get("education"):
        story.append(Paragraph("Pendidikan", section))
        for edu in profile["education"]:
            line = f"<b>{esc(edu.get('jenjang'))}</b> — {esc(edu.get('institusi'))}"
            extras = []
            if edu.get("jurusan"):
                extras.append(esc(edu["jurusan"]))
            if edu.get("tahun_lulus"):
                extras.append(str(edu["tahun_lulus"]))
            if edu.get("ipk"):
                extras.append(f"IPK {edu['ipk']}")
            if extras:
                line += " (" + ", ".join(extras) + ")"
            story.append(Paragraph(line, body))

    if profile.get("skills") or profile.get("certifications") or profile.get("languages"):
        story.append(Paragraph("Skill & Sertifikasi", section))
        if profile.get("skills"):
            story.append(Paragraph("<b>Skill:</b> " + esc(", ".join(profile["skills"])), body))
        for cert in profile.get("certifications", []):
            cert_line = esc(cert.get("nama"))
            if cert.get("penerbit"):
                cert_line += f" — {esc(cert['penerbit'])}"
            if cert.get("tahun"):
                cert_line += f" ({cert['tahun']})"
            story.append(Paragraph("• " + cert_line, body))
        if profile.get("languages"):
            langs = ", ".join(
                f"{lang.get('bahasa')} ({lang.get('tingkat')})" for lang in profile["languages"]
            )
            story.append(Paragraph("<b>Bahasa:</b> " + esc(langs), body))

    story.append(Paragraph("Data Penempatan", section))
    placement_rows = []
    readiness_map = {"segera": "Segera", "n_minggu": "n minggu", "belum_tentu": "Belum tentu"}
    if profile.get("readiness"):
        val = readiness_map.get(profile["readiness"], profile["readiness"])
        if profile["readiness"] == "n_minggu" and profile.get("readiness_weeks"):
            val += f" ({profile['readiness_weeks']} minggu)"
        placement_rows.append(("Kesiapan penempatan", val))
    if profile.get("willing_locations"):
        placement_rows.append(("Lokasi yang bersedia", ", ".join(profile["willing_locations"])))
    if profile.get("expected_salary"):
        placement_rows.append(("Ekspektasi gaji", f"Rp{profile['expected_salary']:,.0f}"))
    if profile.get("contract_preference"):
        placement_rows.append(("Preferensi kontrak", profile["contract_preference"]))
    for label, value in placement_rows:
        story.append(Paragraph(f"<b>{label}:</b> {esc(value)}", body))

    footer = branding.footer_text.strip()
    if footer:

        def _footer(canvas, _doc):
            canvas.saveState()
            canvas.setFont("Helvetica", 7.5)
            canvas.setFillColor(colors.HexColor("#777777"))
            canvas.drawCentredString(A4[0] / 2, 10 * mm, footer[:255])
            canvas.restoreState()

        doc.build(story, onFirstPage=_footer, onLaterPages=_footer)
    else:
        doc.build(story)
    return buffer.getvalue()
